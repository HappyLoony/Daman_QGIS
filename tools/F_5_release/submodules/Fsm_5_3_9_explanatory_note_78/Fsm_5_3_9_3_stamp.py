# -*- coding: utf-8 -*-
"""Fsm_5_3_9_3 — Штамп ГОСТ 21.101-2020 (формы 5/6) + расчёт листа X из Y.

Программное построение штампа через python-docx (без шаблонов skeleton.docx).
Размеры взяты из ГОСТ 21.101-2020 Приложение Ж (img-025/026) и Приложение И (img-027).

Состав:
- build_stamp(metadata): создаёт Document с настроенными секциями и 4 колонтитулами
- calc_first_sheet_number(metadata): расчёт первого листа по cover/title_start
- renumber_sheets(doc, first_sheet_num): перенумерация графы 7 в footer'ах после генерации тела

Принцип: всё в коде. Правки штампа — через диалог Claude, не через Word.
См. feedback_docx_generation.md.
"""
from typing import Optional, Dict, Any

from Daman_QGIS.utils import log_info, log_warning


class ExplanatoryNoteStamp:
    """Программное построение штампа ГОСТ 21.101-2020 формы 5 и 6."""

    # Размеры по ГОСТ 21.101-2020 Приложение Ж + И (всё в мм)
    PAGE_WIDTH_CM = 21.0
    PAGE_HEIGHT_CM = 29.7
    # Body padding 5мм от рамки (left/right):
    # margin_left=2.5см → body_left=25мм = 5мм от frame_left (20мм)
    # margin_right=1.0см → body_right=200мм = 5мм от frame_right (205мм)
    # Form 5/6 (anchor=page absolute) остаются на frame edges независимо от margin.
    MARGIN_LEFT_CM = 2.50
    MARGIN_RIGHT_CM = 1.00
    # Шаг 5мм: 5 (рамка) / 10 (header start) / 15 (header end + line) / 20 (body start).
    MARGIN_TOP_CM = 2.00
    MARGIN_BOTTOM_CM = 0.50

    # === UNIFIED TWIPS CONSTANTS (pixel-perfect alignment) ===
    # Все positioned tables (form 5, form 6, side stamp) и pgBorders должны
    # совпадать на FRAME edges by twips, не по mm (округление mm→twips даёт
    # 0.05-0.14мм mismatch которые Word рендерит как видимые offsets).
    TWIPS_PER_MM = 56.6929  # exact (1pt = 20 twips, 1mm = 2.83465pt)

    # Page и margins в twips (соответствуют Cm()-генерируемым pgMar values):
    PAGE_W_TWIPS = 11906     # 210mm
    PAGE_H_TWIPS = 16838     # 297mm
    MARGIN_T_TWIPS = 283     # 5mm (Cm(0.5) → 283 twips)
    MARGIN_B_TWIPS = 283     # 5mm
    MARGIN_L_TWIPS = 1134    # 20mm (Cm(2.0) → 1134)
    MARGIN_R_TWIPS = 283     # 5mm

    # FRAME edges в twips (== pgBorders position при offsetFrom=text space=0):
    FRAME_T_TWIPS = MARGIN_T_TWIPS                       # 283 (5mm от верха)
    FRAME_B_TWIPS = PAGE_H_TWIPS - MARGIN_B_TWIPS        # 16555 (5mm от низа)
    FRAME_L_TWIPS = MARGIN_L_TWIPS                       # 1134 (20mm)
    FRAME_R_TWIPS = PAGE_W_TWIPS - MARGIN_R_TWIPS        # 11623 (5mm от правого)
    FRAME_W_TWIPS = FRAME_R_TWIPS - FRAME_L_TWIPS        # 10489 (185mm)
    FRAME_H_TWIPS = FRAME_B_TWIPS - FRAME_T_TWIPS        # 16272 (287mm)

    # Форма 5 (заглавный лист) — общая ширина 185 мм
    F5_WIDTH_MM = 185
    # Шапка изменений 14-19: Изм/Кол.уч/Лист/№док/Подп/Дата
    F5_CHANGES_COLS_MM = (10, 10, 10, 10, 15, 10)
    F5_DOC_CODE_MM = 70    # графа (1)
    F5_DOC_TITLE_MM = 70   # графа (5)
    F5_STAGE_MM = 15       # графа (6)
    F5_SHEET_MM = 15       # графа (7)
    F5_TOTAL_MM = 20       # графа (8)
    F5_FIO_HEADER_ROWS = 8
    F5_ROW_HEIGHT_MM = 5

    # Форма 6 (последующие листы)
    F6_WIDTH_MM = 185
    F6_CHANGES_COLS_MM = (10, 10, 10, 10, 15, 10)
    F6_DOC_CODE_MM = 110   # графа (1)
    F6_SHEET_MM = 10       # графа (7)

    # Доп. графы поля подшивки (формы 5 и 6) — ширина 25 мм
    SIDE_AGREED_HEIGHT_MM = 15      # «Согласовано» (только форма 5)
    SIDE_REPLACE_HEIGHT_MM = 25     # Взам. инв. №
    SIDE_SIGNATURE_HEIGHT_MM = 40   # Подп. и дата
    SIDE_INV_NUM_HEIGHT_MM = 25     # Инв. № подл.

    # Маппинг ключей метаданных проекта → переменных штампа.
    # Длинные ключи соответствуют именам в таблице GeoPackage _metadata
    # (см. base_metadata_dialog._map_short_keys_to_full и BASE_METADATA_FIELDS).
    METADATA_TO_STAMP_KEYS: Dict[str, str] = {
        'object_full_name': '1_1_object_full_name',  # полное название объекта
        'doc_code': '2_1_code',                      # шифр документа (графа 1)
        'document_kind': '1_6_stage',                # этап разработки (графа 6, в т.ч. ДПТ)
        'organization': '2_3_company',               # организация-разработчик (графа 9)
        'customer_name': '2_5_customer',             # заказчик (графа 27, штриховая)
        'developer_name': '2_11_developer',          # разработчик ФИО (графа 11)
        'checker_name': '2_12_examiner',             # проверяющий ФИО (графа 11)
        'chief_name': '2_15_chief_of_department',    # руководитель отдела ФИО (графа 11)
        'control_name': '2_13_quality_control',      # нормоконтроль ФИО (графа 11)
        'date': '2_2_date',                          # дата выпуска (графа 13)
    }

    DOCUMENT_TITLE = 'Пояснительная записка'

    # Counter для уникальных docPr id (Word требует unique id per document)
    _drawing_id_counter = 0

    @classmethod
    def mm_twips(cls, mm: float) -> int:
        """Convert mm → integer twips (consistent rounding для всех координат)."""
        return int(round(mm * cls.TWIPS_PER_MM))

    def __init__(self, styles=None):
        """
        Args:
            styles: ExplanatoryNoteStyles instance для применения стилей.
                Если None, создастся новый.
        """
        if styles is None:
            from .Fsm_5_3_9_4_styles import ExplanatoryNoteStyles
            styles = ExplanatoryNoteStyles()
        self.styles = styles

    def build_stamp(self, metadata: Dict[str, Any], first_sheet_num: Optional[int] = None):
        """Создать Document с настроенными секциями и заполненными колонтитулами.

        Args:
            metadata: метаданные проекта (см. METADATA_TO_STAMP_KEYS).
            first_sheet_num: номер первого листа пояснительной записки.
                Если None — рассчитывается через calc_first_sheet_number(metadata).

        Returns:
            docx.Document с пустым телом и готовым штампом.
            Графы 7 (Лист) и 8 (Листов) — Word-поля PAGE/NUMPAGES с offset,
            обновляются автоматически при открытии в Word.
        """
        from docx import Document
        from docx.shared import Cm

        doc = Document()
        # КРИТИЧНО: compatibilityMode=15 заставляет Word рендерить positioned tables
        # в modern mode. Без этого Word legacy mode сдвигает borders на 108 twips
        # (~1.9мм) с каждой стороны (= ~3.8мм mismatch с pgBorders). Это и был наш bug.
        # См. [MS-OI29500] §17.4.57; research compass_artifact_wf-8011e7fe.
        self._add_compat_settings(doc)

        section = doc.sections[0]
        section.page_width = Cm(self.PAGE_WIDTH_CM)
        section.page_height = Cm(self.PAGE_HEIGHT_CM)
        section.left_margin = Cm(self.MARGIN_LEFT_CM)
        section.right_margin = Cm(self.MARGIN_RIGHT_CM)
        section.top_margin = Cm(self.MARGIN_TOP_CM)
        section.bottom_margin = Cm(self.MARGIN_BOTTOM_CM)
        # header_distance=10мм. Header text «Чижик_новый» в [10, 15]мм →
        # 5мм padding от рамки top. Body content на 20мм (margin_top).
        # footer_distance=5мм для inline form 6 в default footer:
        # footer bottom edge at page_h - 5mm = 292mm = frame_bottom.
        # Form 5 (positioned, first_page) не зависит от footer_distance.
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(0.5)
        section.different_first_page_header_footer = True

        # Рамка через DrawingML rectangle (Approach D из research compass_artifact).
        # behindDoc=1 → body content renders ПОВЕРХ рамки → не блокирует.
        # Положена в header → автоматически повторяется на каждой странице.
        # Все 4 borders в одной системе координат с positioned tables (anchor=page) →
        # pixel-perfect by design (single coordinate system, не arithmetic of two).
        self._build_frame_drawing(section.first_page_header)
        self._build_frame_drawing(section.header)

        ctx = self._build_stamp_context(metadata)

        if first_sheet_num is None:
            first_sheet_num = self.calc_first_sheet_number(metadata)
        offset = first_sheet_num - 1

        self._build_central_header(section.first_page_header, ctx)
        self._build_central_header(section.header, ctx)
        self._build_side_stamp(section.first_page_header)
        self._build_side_stamp(section.header)
        self._build_footer_form5(section.first_page_footer, ctx, offset=offset)
        self._build_footer_form6(section.footer, ctx, offset=offset)

        log_info(
            f"Fsm_5_3_9_3 (build_stamp): штамп создан, "
            f"first_sheet_num={first_sheet_num}, offset={offset}"
        )
        return doc

    def _build_side_stamp(self, header) -> None:
        """Боковой штамп слева (positioned table в headerе).

        Параметры из эталона ЧИЖИК (header2.xml Tbl0):
        - 3 строки × 2 колонки
        - Колонки: col0=5 мм (label вертикально), col1=7 мм (для подписи, пустая)
        - Высоты строк: 25 / 35 / 25 мм (всего 85 мм)
        - Тексты в col0: «Взам. инв. №», «Подпись и дата», «Инв. № подл.»
        - textDirection btLr (вертикальный текст, читается снизу вверх)
        - Position: x=0.80 см от левого края → правый край = 2.00 см = левая
          граница pgBorders. Side stamp «соединён» с рамкой левой гранью.
          y=20.69 см от верха страницы. Низ side stamp = 20.69+8.5 = 29.19 см
          = bottom рамки (соединение с нижней рамкой).

        Видим на обеих страницах (first_page и default header).
        """
        from docx.shared import Mm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        labels = ('Взам. инв. №', 'Подпись и дата', 'Инв. № подл.')
        row_heights_mm = (25, 35, 25)
        col_widths_mm = (5, 7)

        table = header.add_table(rows=3, cols=2, width=Mm(sum(col_widths_mm)))
        # Все 6 границ
        self._set_table_borders(table, top=True, left=True, right=True,
                                bottom=True, insideH=True, insideV=True, sz=4)

        # Fixed widths (Word без этого усредняет column widths)
        self._apply_fixed_widths(table, col_widths_mm)

        # Высоты строк
        for ri, h_mm in enumerate(row_heights_mm):
            tr = table.rows[ri]._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            trHeight = OxmlElement('w:trHeight')
            # mm_twips для consistent rounding с FRAME_*_TWIPS calculations
            trHeight.set(qn('w:val'), str(self.mm_twips(h_mm)))
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)

        # Лейблы в col0 с вертикальным текстом btLr
        for ri, label in enumerate(labels):
            cell = table.rows[ri].cells[0]
            self._set_cell_vertical_text(cell, label, size=8, font='Arial')

        # col1 — пустые ячейки для подписей (тоже btLr чтобы место было таким же)
        for ri in range(3):
            cell = table.rows[ri].cells[1]
            self._set_cell_text_direction(cell, 'btLr')

        # PIXEL-PERFECT alignment в twips:
        # right_edge = FRAME_L (1134) by construction → x = FRAME_L - width
        # bottom_edge = FRAME_B (16555) by construction → y = FRAME_B - height
        side_width_twips = sum(self.mm_twips(w) for w in col_widths_mm)
        side_height_twips = sum(self.mm_twips(h) for h in row_heights_mm)
        side_x_twips = self.FRAME_L_TWIPS - side_width_twips
        side_y_twips = self.FRAME_B_TWIPS - side_height_twips
        self._position_table(
            table,
            x_twips=side_x_twips,
            y_twips=side_y_twips,
            width_twips=side_width_twips,
            anchor='page',
        )

    def _set_row_height_exact(self, row, height_mm: float) -> None:
        """Зафиксировать высоту строки таблицы (hRule=exact, не растягивается).

        По ГОСТ 21.101-2020 Приложение Ж формы 5/6 строки штампа имеют
        фиксированные высоты (5 мм для шапки/ФИО, 10 мм для (27)).
        Без hRule="exact" Word растягивает строки по содержимому.

        Args:
            row: docx.table._Row.
            height_mm: высота в мм.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        # Удалить существующие trHeight
        for existing in trPr.findall(qn('w:trHeight')):
            trPr.remove(existing)
        trHeight = OxmlElement('w:trHeight')
        # mm_twips для consistent rounding с FRAME_*_TWIPS (56.6929 exact)
        trHeight.set(qn('w:val'), str(self.mm_twips(height_mm)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

    def _set_cell_text_direction(self, cell, direction: str = 'btLr') -> None:
        """Задать направление текста ячейки (btLr = вертикально снизу вверх)."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tcPr = cell._tc.get_or_add_tcPr()
        # Удалить существующий
        existing = tcPr.find(qn('w:textDirection'))
        if existing is not None:
            tcPr.remove(existing)
        textDir = OxmlElement('w:textDirection')
        textDir.set(qn('w:val'), direction)
        tcPr.append(textDir)

    def _set_cell_vertical_text(self, cell, text: str, size: int = 8,
                                font: str = 'Arial', bold: bool = False) -> None:
        """Записать вертикальный текст (btLr) в ячейку."""
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Очистить
        cell.text = ''
        # Установить direction
        self._set_cell_text_direction(cell, 'btLr')
        # Добавить run
        p = cell.paragraphs[0]
        run = p.add_run(text)
        self.styles.set_run_style(run, font=font, size_pt=size, bold=bold)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _build_central_header(self, header, ctx: Dict[str, str], add_top_border: bool = True) -> None:
        """Центральная строка над основной надписью.

        Содержит ТОЛЬКО полное название объекта из 1_1_object_full_name —
        без обёрток («Внесение изменений в проект межевания...») и без кавычек.
        Заполняется как есть из метаданных. TNR 12 курсив, по центру.

        После текста добавляется горизонтальная inline-таблица шириной 17.75 см
        с bottom border — это верхняя горизонтальная линия рамки страницы
        (в эталоне ЧИЖИК это header2.xml Tbl1).

        Args:
            header: section.header или section.first_page_header.
            ctx: контекст штампа.
            add_top_border: добавить горизонтальную линию (True для обоих headers).
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Mm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Очищаем существующий параграф
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = ''
        for run in list(p.runs):
            run._element.getparent().remove(run._element)

        # Padding от рамки 5мм даётся через section margins (margin_left=2.5см,
        # margin_right=1.0см) → body width [25мм, 200мм]. Indent не нужен.

        # EXACT line spacing 5мм через w:spacing w:line w:lineRule="exact".
        # 5мм = 283 twips. Header content height = exactly 5мм → header end at 15мм
        # (= header_distance 10мм + 5мм). Шаг 5мм по ГОСТ.
        pPr = p._p.get_or_add_pPr()
        # Удалить existing spacing если есть (может быть от style)
        for existing_spacing in pPr.findall(qn('w:spacing')):
            pPr.remove(existing_spacing)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(self.mm_twips(5)))  # 283 twips = 5мм
        spacing.set(qn('w:lineRule'), 'exact')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        pPr.append(spacing)

        # Bottom border paragraph = подчёркивание под текстом, на грани 15мм.
        # Через w:pBdr w:bottom — это объединяет text и линию в один элемент
        # (вместо inline-table) → header height = exactly 5мм.
        if add_top_border:
            # Удалить existing pBdr
            existing_bdr = pPr.find(qn('w:pBdr'))
            if existing_bdr is not None:
                pPr.remove(existing_bdr)
            pBdr = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '4')      # 0.5pt
            bottom_border.set(qn('w:space'), '0')   # без gap между text и border
            bottom_border.set(qn('w:color'), 'auto')
            pBdr.append(bottom_border)
            pPr.append(pBdr)

        text = ctx.get('object_full_name', '')
        run = p.add_run(text)
        self.styles.set_run_style(run, italic=True)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _add_compat_settings(self, doc) -> None:
        """Установить compatibility settings в word/settings.xml.

        КРИТИЧНО: python-docx Document() по умолчанию создаёт settings.xml
        с compatibilityMode val=14 (Word 2010 legacy). Простой append val=15
        даёт ДУБЛИКАТ — Word использует первую запись (val=14) и игнорирует нашу.
        Поэтому REMOVE existing compatibilityMode entries, потом append val=15.

        - compatibilityMode=15 — Word 2013+ table-edge semantics. Без этого Word
          legacy mode сдвигает positioned table borders на 108 twips (~1.9мм)
          с каждой стороны = ~3.8мм mismatch с pgBorders.
        - doNotVertAlignCellWithSp — отключает vertical align cells с positioned shapes.
        - doNotBreakConstrainedForcedTable — стабилизирует positioned tables на page break.

        Источник: research compass_artifact_wf-8011e7fe / [MS-OI29500] §17.4.57.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        settings = doc.settings.element
        compat = settings.find(qn('w:compat'))
        if compat is None:
            compat = OxmlElement('w:compat')
            settings.append(compat)

        # Удалить ВСЕ existing compatibilityMode entries (избежать дубликатов)
        for cs in list(compat.findall(qn('w:compatSetting'))):
            if cs.get(qn('w:name')) == 'compatibilityMode':
                compat.remove(cs)

        # Append fresh compatibilityMode=15 (теперь единственный)
        cs = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'), '15')
        compat.append(cs)

        # Дополнительные флаги (idempotent — проверяем перед добавлением)
        for flag_name in ('doNotVertAlignCellWithSp', 'doNotBreakConstrainedForcedTable'):
            existing_flag = compat.find(qn(f'w:{flag_name}'))
            if existing_flag is None:
                compat.append(OxmlElement(f'w:{flag_name}'))

        # NB: НЕ используем <w:updateFields val="true"/> — Word показывает
        # диалог "Документ содержит поля ссылающиеся на другие файлы".
        # Вместо этого помечаем PAGE/NUMPAGES individually через w:dirty="true"
        # на их fldChar elements (см. _set_cell_field).

    def _apply_fixed_widths(self, table, widths_mm) -> None:
        """Установить fixed table layout с explicit column widths (no autofit).

        Без этого Word с tblLayout=auto использует средние значения widths и
        игнорирует custom column widths (tcW per cell). Required для form 5/6
        и side stamp где column widths critical для ГОСТ layout.

        - tblLayout=fixed
        - tblW=dxa с total width
        - tblGrid с explicit gridCol per column
        - tcW per cell (для compatibility)
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Mm

        widths_twips = [self.mm_twips(w) for w in widths_mm]
        total_twips = sum(widths_twips)

        # Per-cell widths
        for i, w_mm in enumerate(widths_mm):
            for cell in table.columns[i].cells:
                cell.width = Mm(w_mm)

        tbl = table._tbl
        tblPr = tbl.tblPr

        # tblLayout=fixed
        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # tblW=dxa explicit total
        existing_w = tblPr.find(qn('w:tblW'))
        if existing_w is not None:
            tblPr.remove(existing_w)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # tblGrid с explicit gridCol per column
        existing_grid = tbl.find(qn('w:tblGrid'))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        tblGrid = OxmlElement('w:tblGrid')
        for w_twips in widths_twips:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(w_twips))
            tblGrid.append(gridCol)
        tblPr.addnext(tblGrid)

    def _suppress_row_bottom_border(self, row) -> None:
        """Убрать bottom border у всех cells в row (через w:tcBorders w:bottom val=nil).

        Используется для last row form 5/6 чтобы их собственный bottom border не
        конкурировал с pgBorders bottom (избежать двойной линии 1.9мм apart).
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            # Удалить existing bottom если есть, append nil
            existing = tcBorders.find(qn('w:bottom'))
            if existing is not None:
                tcBorders.remove(existing)
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            tcBorders.append(bottom)

    def _build_frame_outer(self, header, sz: int = 4) -> None:
        """Внешняя рамка ГОСТ 5/5/5/20 как positioned table в header.

        Approach E из research compass_artifact_wf-8011e7fe — заменяет pgBorders
        на positioned table, чтобы все элементы (frame, side stamp, form 5/6)
        были в единой системе координат (anchor=page) → pixel-perfect alignment
        by mathematical necessity.

        Положена в header → автоматически повторяется на каждой странице.
        Body content рендерится поверх (positioned tables в header не вызывают
        body content avoidance, в отличие от body positioned).

        Геометрия:
        - x=20мм, y=5мм (frame_top_left)
        - Width=185мм, height=287мм (frame inner area полная)
        - Borders: top/left/right=single, bottom=nil
        - Bottom рамки даёт собственный border form 5/6 (на 292мм = совпадает)
        """
        from docx.shared import Mm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # 1×1 table 185×287мм
        table = header.add_table(rows=1, cols=1, width=Mm(185))

        # Row height = 287мм exact (через trHeight)
        tr = table.rows[0]._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(round(287 * 56.69))))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        # Borders: top/left/right=single 0.5pt, bottom=nil
        self._set_table_borders(table, top=True, left=True, right=True,
                                bottom=False, insideH=False, insideV=False, sz=sz)

        # Положить anchor=page: x=20мм=frame_left, y=5мм=frame_top
        self._position_table(table, x_cm=2.0, y_cm=0.5, width_cm=18.5, anchor='page')

    def _build_frame_drawing(self, header) -> None:
        """Рамка ГОСТ 5/5/5/20 через DrawingML rectangle (Approach D из research).

        Преимущества vs pgBorders / positioned table:
        - DrawingML rectangle support behindDoc=1 → body content renders ПОВЕРХ →
          frame не блокирует body (в отличие от positioned table в header).
        - Положена в header → автоматически повторяется на каждой странице.
        - Coordinates anchored to page в EMU → pixel-perfect alignment с positioned
          tables (anchor=page, twips). 1 twip = 635 EMU, 1mm = 36000 EMU.

        Геометрия (соответствует FRAME_*_TWIPS):
        - x = 20 мм = 720000 EMU
        - y = 5 мм  = 180000 EMU
        - cx = 185 мм = 6660000 EMU (frame width)
        - cy = 287 мм = 10332000 EMU (frame height)
        - Line width 6350 EMU = 0.5 pt (по ГОСТ 2.303 thin solid)
        """
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement, parse_xml

        # 1 mm = 36000 EMU
        x_emu = 20 * 36000      # 720000
        y_emu = 5 * 36000       # 180000
        cx_emu = 185 * 36000    # 6660000
        cy_emu = 287 * 36000    # 10332000
        line_w_emu = 6350       # 0.5pt = 6350 EMU

        # Уникальный id per Drawing — Word требует unique docPr ids per document.
        # Без этого Word warning "содержимое не удалось прочитать" при открытии.
        ExplanatoryNoteStamp._drawing_id_counter += 1
        drawing_id = ExplanatoryNoteStamp._drawing_id_counter

        # Build DrawingML rectangle через raw XML (proще чем OxmlElement chain
        # из-за множества namespaces: w, mc, wp, a, wps).
        drawing_xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
  <w:r>
    <mc:AlternateContent>
      <mc:Choice Requires="wps">
        <w:drawing>
          <wp:anchor distT="0" distB="0" distL="0" distR="0"
                     simplePos="0" relativeHeight="1" behindDoc="1"
                     locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
            <wp:extent cx="{cx_emu}" cy="{cy_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{drawing_id}" name="GostFrame{drawing_id}"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr/>
                  <wps:spPr>
                    <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx_emu}" cy="{cy_emu}"/></a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                    <a:noFill/>
                    <a:ln w="{line_w_emu}"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>
                  </wps:spPr>
                  <wps:bodyPr/>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>
      </mc:Choice>
    </mc:AlternateContent>
  </w:r>
</w:p>'''

        # Parse и append к header element
        new_p = parse_xml(drawing_xml)
        header._element.append(new_p)

    def _set_page_borders(self, section, sz: int = 4) -> None:
        """Рамка страницы по ГОСТ 21.101-2020 (offsetFrom=text, space по сторонам).

        Margins должны быть установлены ДО вызова: top=1.5cm, bottom=1.5cm,
        left=2.0cm, right=0.5cm.

        offsetFrom=text + space=0 на всех сторонах → рамка ровно по margins.
        Margins = ГОСТ 5/5/5/20 → рамка на 5/5/5/20 мм от края страницы.
        Body без padding (body_left = margin_left = frame_left = 20мм) →
        body align с form 5/6 по ширине.

        sz=4 (Word units) = 0.5pt — single border по ГОСТ 2.303.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        sectPr = section._sectPr
        existing = sectPr.find(qn('w:pgBorders'))
        if existing is not None:
            sectPr.remove(existing)

        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'text')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(sz))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
        sectPr.append(pgBorders)

    def _set_table_borders(self, table, top: bool = True, left: bool = True,
                           right: bool = True, bottom: bool = True,
                           insideH: bool = True, insideV: bool = True,
                           sz: int = 4) -> None:
        """Гибкая установка границ таблицы (по сторонам).

        Используется для горизонтальных линий-разделителей и форм 5/6
        где нужны конкретные грани (не все 6).
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl = table._tbl
        tblPr = tbl.tblPr
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)

        tblBorders = OxmlElement('w:tblBorders')
        flags = {
            'top': top, 'left': left, 'bottom': bottom, 'right': right,
            'insideH': insideH, 'insideV': insideV,
        }
        for name, enabled in flags.items():
            border = OxmlElement(f'w:{name}')
            if enabled:
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(sz))
                border.set(qn('w:color'), 'auto')
            else:
                border.set(qn('w:val'), 'nil')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def _position_table(self, table, x_cm: float = None, y_cm: float = None,
                        width_cm: float = None, anchor: str = 'page',
                        x_twips: int = None, y_twips: int = None,
                        width_twips: int = None,
                        top_from_text_twips: int = 0) -> None:
        """Позиционировать таблицу абсолютно (floating table).

        Координаты можно передавать в cm (x_cm/y_cm/width_cm) ИЛИ raw twips
        (x_twips/y_twips/width_twips). Twips приоритетен — для pixel-perfect
        alignment без mm→twips rounding loss.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Twips приоритетен. Cm fallback (1 cm ≈ 567 twips).
        if x_twips is None:
            x_twips = int(round(x_cm * 567))
        if y_twips is None:
            y_twips = int(round(y_cm * 567))

        tbl = table._tbl
        tblPr = tbl.tblPr

        # Удалить существующее tblpPr если есть
        existing = tblPr.find(qn('w:tblpPr'))
        if existing is not None:
            tblPr.remove(existing)

        tblpPr = OxmlElement('w:tblpPr')
        # leftFromText/rightFromText/bottomFromText=0 — без horizontal/bottom wrap buffer.
        # topFromText контролируется параметром (selectively для form 5 — buffer 5мм
        # выше штампа на стр 1; для side stamp / form 6 = 0 для нормального layout).
        tblpPr.set(qn('w:leftFromText'), '0')
        tblpPr.set(qn('w:rightFromText'), '0')
        tblpPr.set(qn('w:topFromText'), str(top_from_text_twips))
        tblpPr.set(qn('w:bottomFromText'), '0')
        # vertAnchor/horzAnchor explicit. Per [MS-OI29500] §17.4.57c — Word default
        # vertAnchor=margin / horzAnchor=text. Для form 5/6 'margin' нужен чтобы
        # stamp coordinates были в той же reference что и pgBorders (offsetFrom=text)
        # → pixel-perfect alignment by same origin, not by arithmetic.
        tblpPr.set(qn('w:vertAnchor'), anchor)
        tblpPr.set(qn('w:horzAnchor'), anchor)
        tblpPr.set(qn('w:tblpX'), str(x_twips))
        tblpPr.set(qn('w:tblpY'), str(y_twips))
        tblPr.append(tblpPr)

        # tblOverlap=never — стабилизирует positioning, блокирует overlap
        tblOverlap = OxmlElement('w:tblOverlap')
        tblOverlap.set(qn('w:val'), 'never')
        tblPr.append(tblOverlap)

        # tblStyle="" — defeats inheritance из Normal Table (108 twips cell margin)
        tblStyle = OxmlElement('w:tblStyle')
        tblStyle.set(qn('w:val'), '')
        tblPr.append(tblStyle)

        # tblInd=0 — explicit, defeats inheritance
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '0')
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

        # tblLayout=fixed — критично, без этого autofit может сжать колонки
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # tblCellMar=0 со всех сторон — устраняет default 108 twips cell margins
        # которые в legacy mode сдвигают table borders на ~1.9мм per side.
        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'), '0')
            mar.set(qn('w:type'), 'dxa')
            tblCellMar.append(mar)
        tblPr.append(tblCellMar)

        if width_cm is not None:
            from docx.shared import Cm
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(int(round(width_cm * 567))))
            tblW.set(qn('w:type'), 'dxa')
        elif width_twips is not None:
            from docx.shared import Cm
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is None:
                tblW = OxmlElement('w:tblW')
                tblPr.append(tblW)
            tblW.set(qn('w:w'), str(width_twips))
            tblW.set(qn('w:type'), 'dxa')

    def _build_footer_form5(self, footer, ctx: Dict[str, str], offset: int = 0) -> None:
        """Footer формы 5 (заглавный лист текстового документа, ГОСТ 21.101-2020).

        Стандартные колонки шапки изменений (10/10/10/10/15/10 = 65 мм) +
        правая часть (70/15/15/20 = 120 мм). Total = 185 мм.

        Шапка ФИО (графы 10-13): Разраб./Проверил/Рук.отдела/Н.контр.
        Размещены в строках 3-6. ФИО merged через col 1-3 (30 мм wide).
        """
        from docx.shared import Mm

        self._clear_container(footer)
        footer.add_paragraph()

        # Колонки формы 5 строго по ГОСТ 21.101-2020 Приложение Ж:
        # шапка изменений 10/10/10/10/15/10 = 65 мм
        changes_cols_mm = self.F5_CHANGES_COLS_MM  # (10, 10, 10, 10, 15, 10) сумма 65
        changes_labels = ('Изм.', 'Кол.уч', 'Лист', '№ док.', 'Подп.', 'Дата')

        total_cols = 10
        col_widths_mm = list(changes_cols_mm) + [
            self.F5_DOC_TITLE_MM,  # 70
            self.F5_STAGE_MM,      # 15
            self.F5_SHEET_MM,      # 15
            self.F5_TOTAL_MM,      # 20
        ]
        # сумма: 65 + 70 + 50 = 185 мм ✓

        # 7 rows по ГОСТ Р 21.101-2020 Annex Ж форма 5: 185×40мм main stamp.
        # Row 0: (27) Заказчик 10мм (опц., dashed top border indicator);
        # Rows 1-2: шапка изменений (labels + values);
        # Rows 3-6: 4 ФИО подписантов (Разраб./Проверил/Рук.отдела/Н.контр.).
        # Heights: 10 + 6×5 = 40мм.
        table = footer.add_table(rows=7, cols=total_cols, width=Mm(self.F5_WIDTH_MM))
        # Borders top/left/right=single, bottom=nil (pgBorders bottom закрывает),
        # insideH/insideV=single для разделения rows и columns внутри штампа.
        self._set_table_borders(table, top=True, left=True, right=True,
                                bottom=False, insideH=True, insideV=True)
        # Fixed widths (Word без этого усредняет column widths)
        self._apply_fixed_widths(table, col_widths_mm)

        row_heights_mm = (10, 5, 5, 5, 5, 5, 5)  # сумма 40мм
        for ri, h_mm in enumerate(row_heights_mm):
            self._set_row_height_exact(table.rows[ri], h_mm)

        # Row 0: графа (27) Заказчик cells 6-9 merged (120мм × 10мм)
        customer = ctx.get('customer_name', '').strip()
        if customer:
            merged = table.rows[0].cells[6].merge(table.rows[0].cells[9])
            self._set_cell_text(merged, customer, size=10, italic=True, align='center')

        # Row 1: лейблы шапки изменений Изм/Кол.уч/Лист/№док/Подп/Дата
        row1 = table.rows[1]
        for i, label in enumerate(changes_labels):
            self._set_cell_text(row1.cells[i], label, size=8, align='center')
        # Right portion row 1: (5) название cell 6 + Стадия/Лист/Листов labels cells 7-9
        self._set_cell_text(row1.cells[6], ctx.get('document_title', ''),
                            size=10, align='center')
        self._set_cell_text(row1.cells[7], 'Стадия', size=8, align='center')
        self._set_cell_text(row1.cells[8], 'Лист', size=8, align='center')
        self._set_cell_text(row1.cells[9], 'Листов', size=8, align='center')

        # Row 2: changes values (пустые) + (1) шифр + Стадия value / PAGE / NUMPAGES
        row2 = table.rows[2]
        self._set_cell_text(row2.cells[6], ctx.get('doc_code', ''),
                            size=14, bold=True, align='center')
        self._set_cell_text(row2.cells[7], ctx.get('document_kind', ''),
                            size=10, align='center')
        self._set_cell_field(row2.cells[8], f'= PAGE + {offset}', size=10, align='center')
        self._set_cell_field(row2.cells[9], f'= NUMPAGES + {offset}', size=10, align='center')

        # Rows 3-6: ФИО подписантов (графы 10-13) по ГОСТ.
        # col 0 = должность, cells 1-3 merged = ФИО (30мм), col 4 = подпись, col 5 = дата.
        fio_rows = (
            ('Разраб.',    ctx.get('developer_name', '')),
            ('Проверил',   ctx.get('checker_name', '')),
            ('Рук.отдела', ctx.get('chief_name', '')),
            ('Н.контр.',   ctx.get('control_name', '')),
        )
        date_str = ctx.get('date', '')
        for i, (role, name) in enumerate(fio_rows, start=3):
            row = table.rows[i]
            self._set_cell_text(row.cells[0], role, size=8, align='center')
            fio_cell = row.cells[1].merge(row.cells[3])
            self._set_cell_text(fio_cell, name, size=9, align='center')
            self._set_cell_text(row.cells[5], date_str, size=8, align='center')

        # Right side (col 6): (5) название row 3 + (9) Орг row 5
        self._set_cell_text(table.rows[3].cells[6], ctx.get('document_title', ''),
                            size=12, align='center')
        self._set_cell_text(table.rows[5].cells[6], ctx.get('organization', ''),
                            size=12, align='center')

        # Suppress own bottom border у last row — рамка закрывает.
        self._suppress_row_bottom_border(table.rows[-1])

        # КРИТИЧНО: form 5 INLINE (без tblpPr) — не positioned floating frame.
        # Floating frame в footer вызывает body wrap-around → huge gap между
        # header_table и body_table на стр 1 (research compass_artifact:
        # "positioned tables в footers cause body content avoidance gaps").
        #
        # Inline form 5 в first_page_footer container, с footer_distance=5mm:
        # form 5 bottom = page_h - 5мм = 292мм = frame_bottom.
        # Form 5 height 40мм → top at 252мм. Body region [20, 252]мм без gap.
        #
        # tblInd=-5мм: form 5 extends 5мм левее margin_left=25мм → starts at 20мм
        # = frame_left. Width 185мм → ends at 205мм = frame_right.
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tblPr = table._tbl.tblPr
        existing_ind = tblPr.find(qn('w:tblInd'))
        if existing_ind is not None:
            tblPr.remove(existing_ind)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(-self.mm_twips(5)))  # -283 twips = -5мм
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    def _build_footer_form6(self, footer, ctx: Dict[str, str], offset: int = 0) -> None:
        """Footer формы 6 (последующие листы) по ГОСТ 21.101-2020 Приложение Ж.

        Структура 3×5=15мм:
        - Row 0 (5мм): пустая (резерв над шапкой изменений)
        - Row 1 (5мм): labels Изм./Кол.уч/Лист/№док/Подп/Дата + (1) шифр + Лист label
        - Row 2 (5мм): values изменений (пустые) + (1) шифр (cont) + (7) PAGE field
        """
        from docx.shared import Mm

        self._clear_container(footer)
        footer.add_paragraph()

        col_widths = list(self.F6_CHANGES_COLS_MM) + [self.F6_DOC_CODE_MM, self.F6_SHEET_MM]
        # 65 + 110 + 10 = 185 ✓
        changes_labels = ('Изм.', 'Кол.уч', 'Лист', '№ док.', 'Подп.', 'Дата')

        table = footer.add_table(rows=3, cols=8, width=Mm(self.F6_WIDTH_MM))
        # Borders top/left/right=single, bottom=nil (pgBorders bottom закрывает).
        self._set_table_borders(table, top=True, left=True, right=True,
                                bottom=False, insideH=True, insideV=True)
        # Fixed widths (Word без этого усредняет column widths)
        self._apply_fixed_widths(table, col_widths)

        # Высоты строк: 3 × 5 мм = 15 мм total (по ГОСТ 21.101-2020 форма 6)
        for row in table.rows:
            self._set_row_height_exact(row, 5)

        # Row 0: пустая (резерв над шапкой изменений), но col 6 = (1) шифр (top part),
        # col 7 = пусто. Шифр merged across rows 0-1 (visually одна ячейка).
        # Для простоты не merge — заполняем col 6 шифром тут же.
        row0 = table.rows[0]
        self._set_cell_text(row0.cells[6], ctx.get('doc_code', ''),
                            size=14, bold=True, align='center')

        # Row 1: лейблы шапки + (1) шифр (continued) + (7) Лист label
        row1 = table.rows[1]
        for i, label in enumerate(changes_labels):
            self._set_cell_text(row1.cells[i], label, size=8, align='center')
        self._set_cell_text(row1.cells[7], 'Лист', size=8, align='center')

        # Row 2: значения изменений (пустые) + (7) графа = PAGE
        row2 = table.rows[2]
        self._set_cell_field(row2.cells[7], f'= PAGE + {offset}', size=10, align='center')

        # Suppress own bottom border у last row — рамка закрывает.
        self._suppress_row_bottom_border(table.rows[-1])

        # КРИТИЧНО: form 6 INLINE (без tblpPr) — не positioned. Иначе Word
        # рассматривает её как floating frame с body wrap, что вызывает huge gap
        # на стр 2+ при tblHeader+cantSplit на T1 row 0 (research compass_artifact).
        # Inline footer table рендерится в footer container естественно.
        #
        # Position via section.footer_distance: footer bottom = page_h - 5mm = 292mm.
        # Form 6 height = 15mm → form 6 [277, 292]mm = aligned с frame bottom.
        #
        # tblInd=-5mm: table extends 5mm к левее margin_left=25mm → starts at 20mm
        # = frame_left. Width 185mm → ends at 205mm = frame_right.
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tblPr = table._tbl.tblPr
        existing_ind = tblPr.find(qn('w:tblInd'))
        if existing_ind is not None:
            tblPr.remove(existing_ind)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(-self.mm_twips(5)))  # -283 twips = -5mm
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    def _set_cell_text(self, cell, text: str, bold: bool = False, italic: bool = False,
                       size: int = 10, align: str = 'left') -> None:
        """Записать текст в ячейку с компактным форматированием (для штампа).

        line_spacing=1.0, space_before=0, space_after=0 — иначе ячейки штампа
        раздуваются Word default'ом 8-10pt после абзаца.
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt

        cell.text = ''  # очистка
        p = cell.paragraphs[0]
        # Компактный paragraph_format (по ГОСТ 21.101-2020 формы 5/6)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        run = p.add_run(text)
        self.styles.set_run_style(run, size_pt=size, bold=bold, italic=italic)
        if align == 'center':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == 'right':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def _clear_container(self, container) -> None:
        """Удалить все paragraphs и tables в header/footer."""
        for p in list(container.paragraphs):
            elem = p._element
            if elem.getparent() is not None:
                elem.getparent().remove(elem)
        for t in list(container.tables):
            elem = t._element
            if elem.getparent() is not None:
                elem.getparent().remove(elem)

    def _build_stamp_context(self, metadata: Dict[str, Any]) -> Dict[str, str]:
        """Собрать значения для штампа. Пустые → '' + warning."""
        ctx = {'document_title': self.DOCUMENT_TITLE}
        for stamp_key, meta_key in self.METADATA_TO_STAMP_KEYS.items():
            value = metadata.get(meta_key, '')
            if not value:
                log_warning(
                    f"Fsm_5_3_9_3 (_build_stamp_context): "
                    f"метаданные {meta_key!r} пустые, штамп: {stamp_key}=''"
                )
            ctx[stamp_key] = str(value or '')
        return ctx

    def calc_first_sheet_number(self, metadata: Dict[str, Any]) -> int:
        """Расчёт номера первого листа пояснительной.

        cover='Наша'      -> 4 (1=обложка, 2=титул, 3=содержание, 4=текст)
        cover='заказчика' -> int(metadata['2_9_title_start'])
        иначе/ошибка      -> 1 + log_warning
        """
        cover = (metadata.get('2_8_cover') or '').strip().lower()

        if cover == 'наша':
            return 4

        if cover in ('заказчика', 'заказчик'):
            title_start = metadata.get('2_9_title_start', '')
            try:
                return int(str(title_start).strip())
            except (ValueError, TypeError):
                log_warning(
                    f"Fsm_5_3_9_3 (calc_first_sheet_number): "
                    f"некорректный 2_9_title_start={title_start!r}, возвращаю 1"
                )
                return 1

        log_warning(
            f"Fsm_5_3_9_3 (calc_first_sheet_number): "
            f"неизвестный 2_8_cover={cover!r}, возвращаю 1"
        )
        return 1

    def _set_cell_field(self, cell, formula: str, size: int = 10, align: str = 'center') -> None:
        """Записать в ячейку Word-поле с вычисляемой формулой типа '= PAGE + 3'.

        Word трактует `= PAGE` в fldSimple как ссылку на закладку 'PAGE', а не как
        системное поле. Поэтому используем nested fldChar: внешнее поле = формула,
        внутреннее = PAGE/NUMPAGES.

        Поддерживаются формулы:
        - '= PAGE + N'      — текущая страница + offset
        - '= NUMPAGES + N'  — всего страниц + offset
        - 'PAGE' / 'NUMPAGES' (без =) — простое поле без формулы

        Структура OOXML для '= PAGE + 3':
            <w:r><w:fldChar w:fldCharType="begin"/></w:r>
            <w:r><w:instrText> = </w:instrText></w:r>
              <w:r><w:fldChar w:fldCharType="begin"/></w:r>
              <w:r><w:instrText> PAGE </w:instrText></w:r>
              <w:r><w:fldChar w:fldCharType="separate"/></w:r>
              <w:r><w:t>1</w:t></w:r>           <!-- кэш внутреннего -->
              <w:r><w:fldChar w:fldCharType="end"/></w:r>
            <w:r><w:instrText> + 3 </w:instrText></w:r>
            <w:r><w:fldChar w:fldCharType="separate"/></w:r>
            <w:r><w:t>1</w:t></w:r>              <!-- кэш внешнего -->
            <w:r><w:fldChar w:fldCharType="end"/></w:r>
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        cell.text = ''
        p = cell.paragraphs[0]
        if align == 'center':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == 'right':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Парсим формулу: ищем PAGE или NUMPAGES + опц. offset
        m = re.match(r'^\s*=\s*(PAGE|NUMPAGES)\s*\+\s*(\d+)\s*$', formula, re.IGNORECASE)
        is_formula = m is not None
        if is_formula:
            inner_field = m.group(1).upper()  # PAGE или NUMPAGES
            offset = int(m.group(2))
        else:
            # Простое поле без формулы: '= PAGE + 0' интерпретируем как просто PAGE
            simple_m = re.match(r'^\s*(?:=\s*)?(PAGE|NUMPAGES)\s*$', formula, re.IGNORECASE)
            if simple_m:
                inner_field = simple_m.group(1).upper()
                offset = 0
                is_formula = False  # fallback на простое поле без обёртки
            else:
                # Не распознали — просто записываем текст
                run = p.add_run(formula)
                self.styles.set_run_style(run, size_pt=size)
                return

        def _make_run_props():
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), self.styles.FONT_NAME)
            rFonts.set(qn('w:hAnsi'), self.styles.FONT_NAME)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(size * 2))
            rPr.append(sz)
            return rPr

        def _add_fld_char(parent, char_type: str):
            r = OxmlElement('w:r')
            r.append(_make_run_props())
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), char_type)
            # NB: w:dirty="true" вызывает Word warning о fields linked to other files,
            # хотя это false positive (PAGE/NUMPAGES внутренние). Не используем.
            # Word обновит fields при печати или Ctrl+A → F9 вручную.
            r.append(fc)
            parent.append(r)

        def _add_instr_text(parent, text: str):
            r = OxmlElement('w:r')
            r.append(_make_run_props())
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = text
            r.append(it)
            parent.append(r)

        def _add_cached_value(parent, value: str = '1'):
            r = OxmlElement('w:r')
            r.append(_make_run_props())
            t = OxmlElement('w:t')
            t.text = value
            r.append(t)
            parent.append(r)

        elem = p._p

        if is_formula:
            # Внешнее = formula
            _add_fld_char(elem, 'begin')
            _add_instr_text(elem, ' = ')
            # Вложенное PAGE / NUMPAGES
            _add_fld_char(elem, 'begin')
            _add_instr_text(elem, f' {inner_field} ')
            _add_fld_char(elem, 'separate')
            _add_cached_value(elem, '1')
            _add_fld_char(elem, 'end')
            # Остаток формулы
            _add_instr_text(elem, f' + {offset} \\* MERGEFORMAT ')
            _add_fld_char(elem, 'separate')
            _add_cached_value(elem, '1')
            _add_fld_char(elem, 'end')
        else:
            # Простое поле PAGE / NUMPAGES без обёртки = ... +
            _add_fld_char(elem, 'begin')
            _add_instr_text(elem, f' {inner_field} \\* MERGEFORMAT ')
            _add_fld_char(elem, 'separate')
            _add_cached_value(elem, '1')
            _add_fld_char(elem, 'end')

    def renumber_sheets(self, doc, first_sheet_num: int) -> None:
        """Обновить offset в Word-полях PAGE/NUMPAGES после генерации тела.

        Ищет instrText содержащий ' + N \\* MERGEFORMAT ' (наш паттерн от
        _set_cell_field) и заменяет N на first_sheet_num - 1.

        Используется если build_stamp был вызван с другим first_sheet_num,
        и нужно пересчитать. В норме build_stamp(metadata, first_sheet_num)
        сразу выставляет правильный offset.
        """
        from docx.oxml.ns import qn
        import re

        new_offset = first_sheet_num - 1
        # Паттерн: ' + 3 \* MERGEFORMAT ' → ' + 9 \* MERGEFORMAT '
        pattern = re.compile(r'(\s\+\s*)\d+(\s*\\\*\s*MERGEFORMAT)', re.IGNORECASE)

        replaced = 0
        for section in doc.sections:
            for header_or_footer in (
                section.first_page_header, section.header,
                section.first_page_footer, section.footer,
            ):
                for it in header_or_footer._element.iter(qn('w:instrText')):
                    text = it.text or ''
                    new_text, n = pattern.subn(rf'\g<1>{new_offset}\g<2>', text)
                    if n:
                        it.text = new_text
                        replaced += n

        log_info(f"Fsm_5_3_9_3 (renumber_sheets): обновлено полей: {replaced}, "
                 f"new offset={new_offset}")
