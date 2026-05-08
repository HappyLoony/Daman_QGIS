# -*- coding: utf-8 -*-
"""Fsm_5_3_9_2 — Программная сборка тела пояснительной записки через python-docx.

Состав:
- intro paragraph (фиксированный текст, JUSTIFY, TNR 12, 1.5 интервал, отступ 1.25 см)
- 7 подразделов: заголовок-абзац (List Paragraph по эталону) + таблица
  T1 — Перечень образуемых ЗУ (7 колонок)
  T2 — Резервирование/изъятие (4 колонки)
  T3 — Объекты недвижимости (4 колонки)
  T4 — Сервитуты (3 колонки)
  T5 — Координаты ЗУ (3 колонки, группировка + строка S=)
  T6 — те же координаты ЗУ что T5, без строки S= (3 колонки, группировка)
  T7 — ВРИ (3 колонки)

Все таблицы рамкой single 0.5pt по 6 направлениям (ГОСТ 2.303). Шрифт TNR.
"""
from typing import Optional, Dict, Any, List

from Daman_QGIS.utils import log_info, log_warning


MODULE_ID = 'Fsm_5_3_9_2'

INTRO_TEXT = (
    "Перечень и сведения о площади образуемых земельных участков, которые будут "
    "отнесены к территориям общего пользования или имуществу общего пользования, "
    "в том числе в отношении которых предполагаются резервирование и (или) изъятие "
    "для государственных нужд, способы образования земельных участков, виды "
    "разрешенного использования образуемых земельных, представлены ниже."
)

SUBSECTION_TITLES: Dict[str, str] = {
    't1_zu': '1. Перечень образуемых земельных участков.',
    't2_reservation': (
        'Перечень образуемых и существующих земельных участков, в отношении '
        'которых предполагаются их резервирование и (или) изъятие для '
        'государственных нужд'
    ),
    't3_realty': (
        'В границах образуемых и существующих земельных участков, в отношении '
        'которых предполагаются их резервирование и (или) изъятие для '
        'государственных нужд расположены следующие объекты недвижимого '
        'имущества, сведения о которых содержатся в Едином государственном '
        'ресстре недвижимости.'
    ),
    't4_servitude': (
        '2. Перечень кадастровых номеров существующих земельных участков, на '
        'которых линейный объект может быть размещен на условиях сервитута, '
        'публичного сервитута, их адреса или описание местоположения.'
    ),
    't5_coords': '3. Перечень координат характерных точек образуемых земельных участков.',
    't6_gpmt_coords': (
        '4. Сведения о границах территории, применительно к которой осуществляется '
        'подготовка проекта межевания, содержащие перечень координат '
        'характерных точек таких границ в системе координат, используемой для '
        'ведения Единого государственного реестра недвижимости.'
    ),
    't7_vri': (
        '5. Вид разрешенного использования образуемых и существующих земельных '
        'участков, предназначенных для размещения {type_phrase}.'
    ),
}


class ExplanatoryNoteDocBuilder:
    """Программная сборка тела документа через python-docx."""

    INTRO_TEXT = INTRO_TEXT
    SUBSECTION_TITLES = SUBSECTION_TITLES

    # Body width при margin L=2.5 / R=1.0 → 17.5 см. Custom widths per таблицу
    # (заданы пользователем для optimal visual fit).
    T1_WIDTHS_CM = (2.0, 2.0, 3.5, 2.0, 2.5, 2.5, 3.0)  # 7 колонок, сумма 17.5
    T2_WIDTHS_CM = (2.0, 4.0, 8.0, 3.5)                  # 4 колонки, сумма 17.5
    T3_WIDTHS_CM = (4.0, 4.0, 6.5, 3.0)                  # 4 колонки, сумма 17.5
    T4_WIDTHS_CM = (2.0, 4.0, 11.5)                      # 3 колонки, сумма 17.5
    T5_WIDTHS_CM = (4.5, 6.5, 6.5)                       # 3 колонки, сумма 17.5 (T5 + T6)
    T7_WIDTHS_CM = (2.0, 4.5, 11.0)                      # 3 колонки, сумма 17.5

    def __init__(self, styles):
        self.styles = styles

    # ------------------------------------------------------------------
    # Public entrypoint
    # ------------------------------------------------------------------
    def build_body(self, doc, data: Dict[str, Any], metadata: Dict[str, Any],
                   progress_callback=None):
        """Заполнить тело документа: intro + 7 подразделов.

        Args:
            doc: docx.Document с уже отрендеренным штампом и пустым телом.
            data: результат ExplanatoryNoteDataCollector.collect_all().
            metadata: метаданные проекта (для подстановок в заголовки).
            progress_callback: callable(msg, percent) — опц., обновление UI.

        Returns:
            Тот же doc, с заполненным телом.
        """
        log_info(f"{MODULE_ID} (build_body): начало сборки тела")

        def _report(msg: str, percent: int) -> None:
            if progress_callback is None:
                return
            try:
                progress_callback(msg, percent)
            except Exception as e:
                log_warning(f"{MODULE_ID} (progress_callback): {e}")

        self._clear_body(doc)
        self._build_intro(doc, metadata)

        # Прогресс распределён 78..94% (16% общего на сборку 7 таблиц)
        _report('Тело: T1 — Перечень образуемых ЗУ', 80)
        self._build_subsection_title(doc, self.SUBSECTION_TITLES['t1_zu'])
        self._build_table_t1(doc, data.get('t1_zu', []))
        self._add_after_table_spacer(doc)

        _report('Тело: T2 — Резервирование/изъятие', 82)
        self._build_subsection_title(doc, self.SUBSECTION_TITLES['t2_reservation'])
        self._build_table_t2(doc, data.get('t2_reservation', []))
        self._add_after_table_spacer(doc)

        _report('Тело: T3 — Объекты недвижимости', 84)
        self._build_subsection_title(doc, self.SUBSECTION_TITLES['t3_realty'])
        self._build_table_t3(doc, data.get('t3_realty', []))
        self._add_after_table_spacer(doc)

        _report('Тело: T4 — Сервитуты', 86)
        self._build_subsection_title(doc, self.SUBSECTION_TITLES['t4_servitude'])
        self._build_table_t4(doc, data.get('t4_servitude', []))
        self._add_after_table_spacer(doc)

        _report('Тело: T5 — Координаты ЗУ', 88)
        self._build_subsection_title(doc, self.SUBSECTION_TITLES['t5_coords'])
        self._build_coords_table(doc, data.get('t5_coords', []), show_area=True)
        self._add_after_table_spacer(doc)

        _report('Тело: T6 — Сведения о границах территории ПМТ', 90)
        self._build_subsection_title(doc, self.SUBSECTION_TITLES['t6_gpmt_coords'])
        self._build_coords_table(doc, data.get('t6_gpmt_coords', []), show_area=False)
        self._add_after_table_spacer(doc)

        _report('Тело: T7 — ВРИ', 92)
        # T7 заголовок dynamic per object type:
        # - Линейный: "размещения линейного объекта <type_value> значения"
        # - Площадной: "размещения площадного объекта" (без значения)
        object_type = str(metadata.get('1_2_object_type', '') or '').strip().lower()
        type_value = str(metadata.get('1_2_1_object_type_value', '') or '').strip()
        is_linear = 'лин' in object_type
        if is_linear:
            type_phrase = f'линейного объекта {type_value} значения' if type_value \
                          else 'линейного объекта'
        else:
            # Площадной/площадка/etc. — без типа значения и уровня
            type_phrase = 'площадного объекта'
        t7_title = self.SUBSECTION_TITLES['t7_vri'].format(type_phrase=type_phrase)
        self._build_subsection_title(doc, t7_title)
        self._build_table_t7(doc, data.get('t7_vri', []))

        log_info(f"{MODULE_ID} (build_body): тело собрано")
        return doc

    # ------------------------------------------------------------------
    # Body manipulation
    # ------------------------------------------------------------------
    def _clear_body(self, doc) -> None:
        """Удалить все параграфы и таблицы тела, оставить только sectPr.

        sectPr в конце <w:body> хранит настройки секции (страница, поля,
        размеры) — его трогать нельзя, штамп через колонтитулы привязан к
        этой секции.
        """
        from docx.oxml.ns import qn

        body = doc.element.body
        sectPr = body.find(qn('w:sectPr'))
        # Удаляем всё кроме sectPr
        for child in list(body):
            if child is sectPr:
                continue
            body.remove(child)

    def _build_intro(self, doc, metadata: Dict[str, Any]) -> None:
        """Вступительный абзац с фиксированным текстом."""
        p = doc.add_paragraph()
        run = p.add_run(self.INTRO_TEXT)
        self.styles.set_run_style(
            run,
            font=self.styles.FONT_NAME,
            size_pt=self.styles.FONT_SIZE_PT,
        )
        self.styles.set_para_style(
            p,
            justify=True,
            first_line_cm=self.styles.FIRST_LINE_INDENT_CM,
            line_spacing=self.styles.LINE_SPACING,
        )

    def _build_subsection_title(self, doc, text: str) -> None:
        """Заголовок подраздела (без номера, по эталону — List Paragraph).

        В эталоне ЧИЖИК заголовки набраны TNR 12, justify, с отступом первой
        строки 1.25 см (как обычный абзац). Без bold.
        """
        p = doc.add_paragraph()
        run = p.add_run(text)
        self.styles.set_run_style(
            run,
            font=self.styles.FONT_NAME,
            size_pt=self.styles.FONT_SIZE_PT,
        )
        self.styles.set_para_style(
            p,
            justify=True,
            first_line_cm=self.styles.FIRST_LINE_INDENT_CM,
            line_spacing=self.styles.LINE_SPACING,
        )

    # ------------------------------------------------------------------
    # T1 — Перечень образуемых ЗУ (7 колонок)
    # ------------------------------------------------------------------
    def _build_table_t1(self, doc, rows: List[Dict[str, Any]]) -> None:
        headers = (
            'Условный номер',
            'Номера характерных точек',
            'Кадастровый номер исходного земельного участка',
            'Площадь образуемого земельного участка, га',
            'Сведения об отнесении (неотнесении) образуемого земельного участка '
            'к территориям общего пользования или имуществу общего пользования',
            'Способ образования земельного участка',
            'Сведения об отнесении земельного участка к определенной категории '
            'земель или о невозможности отнесения к определенной категории земель',
        )
        header_table, body_table = self._make_table(doc, n_cols=7, widths=self.T1_WIDTHS_CM)
        self._fill_header(header_table, headers)
        self._fill_numbering_row(body_table, n_cols=7)

        if not rows:
            self._append_empty_row(body_table, n_cols=7)
            return

        for row in rows:
            tr = body_table.add_row()
            cells = tr.cells
            self._set_cell_text(cells[0], str(row.get('id', '-')), size=12, align='center')
            self._set_cell_text(cells[1], str(row.get('points', '-')), size=12, align='center')
            self._set_cell_text(cells[2], str(row.get('kn', '-')), size=12, align='center')
            self._set_cell_text(cells[3], str(row.get('area_ha', '-')), size=12, align='center')
            self._set_cell_text(cells[4], str(row.get('common_land', '-')), size=12, align='center')
            self._set_cell_text(cells[5], str(row.get('work_kind', '-')), size=12, align='center')
            self._set_cell_text(cells[6], str(row.get('category', '-')), size=12, align='center')

    # ------------------------------------------------------------------
    # T2 — Резервирование/изъятие (4 колонки)
    # ------------------------------------------------------------------
    def _build_table_t2(self, doc, rows: List[Dict[str, Any]]) -> None:
        headers = (
            'Условный номер',
            'Кадастровый номер существующего земельного участка',
            'Адрес или описание местоположения',
            'Резервирование и (или) изъятие для государственных или '
            'муниципальных нужд',
        )
        header_table, body_table = self._make_table(doc, n_cols=4, widths=self.T2_WIDTHS_CM)
        self._fill_header(header_table, headers)
        self._fill_numbering_row(body_table, n_cols=4)

        if not rows:
            self._append_empty_row(body_table, n_cols=4)
            return

        for row in rows:
            tr = body_table.add_row()
            cells = tr.cells
            self._set_cell_text(cells[0], str(row.get('id', '-')), size=12, align='center')
            self._set_cell_text(cells[1], str(row.get('kn_existing', '-')), size=12, align='center')
            self._set_cell_text(cells[2], str(row.get('address', '-')), size=12, align='left')
            self._set_cell_text(cells[3], str(row.get('reservation_kind', '-')), size=12, align='center')

    # ------------------------------------------------------------------
    # T3 — Объекты недвижимости (4 колонки)
    # ------------------------------------------------------------------
    def _build_table_t3(self, doc, rows: List[Dict[str, Any]]) -> None:
        headers = (
            'Кадастровый номер существующего земельного участка',
            'Кадастровый номер объекта недвижимости',
            'Адрес или описание местоположения объекта недвижимости',
            'Наименование объекта недвижимости',
        )
        header_table, body_table = self._make_table(doc, n_cols=4, widths=self.T3_WIDTHS_CM)
        self._fill_header(header_table, headers)
        self._fill_numbering_row(body_table, n_cols=4)

        if not rows:
            self._append_empty_row(body_table, n_cols=4)
            return

        for row in rows:
            tr = body_table.add_row()
            cells = tr.cells
            self._set_cell_text(cells[0], str(row.get('kn_zu', '-')), size=12, align='center')
            self._set_cell_text(cells[1], str(row.get('kn_oks', '-')), size=12, align='center')
            self._set_cell_text(cells[2], str(row.get('address', '-')), size=12, align='left')
            self._set_cell_text(cells[3], str(row.get('name', '-')), size=12, align='left')

    # ------------------------------------------------------------------
    # T4 — Сервитуты (3 колонки)
    # ------------------------------------------------------------------
    def _build_table_t4(self, doc, rows: List[Dict[str, Any]]) -> None:
        headers = (
            '№ п/п',
            'Кадастровый номер существующего земельного участка',
            'Адрес или описание местоположения',
        )
        header_table, body_table = self._make_table(doc, n_cols=3, widths=self.T4_WIDTHS_CM)
        self._fill_header(header_table, headers)
        self._fill_numbering_row(body_table, n_cols=3)

        if not rows:
            self._append_empty_row(body_table, n_cols=3)
            return

        for row in rows:
            tr = body_table.add_row()
            cells = tr.cells
            self._set_cell_text(cells[0], str(row.get('num', '-')), size=12, align='center')
            self._set_cell_text(cells[1], str(row.get('kn_existing', '-')), size=12, align='center')
            self._set_cell_text(cells[2], str(row.get('address', '-')), size=12, align='left')

    # ------------------------------------------------------------------
    # T5 / T6 — Координаты (3 колонки, группировка)
    # ------------------------------------------------------------------
    def _build_coords_table(self, doc, groups: List[Dict[str, Any]],
                            show_area: bool = True) -> None:
        """3 колонки: Номер точки | X(м) | Y(м).

        Структура:
            [шапка] Номер точки | X(м) | Y(м)
            [нумерация] 1 | 2 | 3
            [group title row] (merge×3, bold) — title группы (внешний контур)
            [points exterior] N | x | y
            [ring subtitle row] (merge×3, bold) — «Внутренний контур M» (опц.)
            [points hole] N | x | y
            ... повтор подзаголовков для последующих holes
            [S= row] (только если show_area=True и area_sqm не пуст)

        Args:
            groups: список групп {title, rings: [{subtitle, points}], area_sqm?}.
                rings[0] — exterior (subtitle=None), rings[1..] — holes
                (subtitle=«Внутренний контур N», формат как у головного title).
            show_area: True для T5 (показывать «S=»), False для T6.
        """
        headers = ('Номер точки', 'X(м)', 'Y(м)')
        header_table, body_table = self._make_table(doc, n_cols=3, widths=self.T5_WIDTHS_CM)
        self._fill_header(header_table, headers)
        self._fill_numbering_row(body_table, n_cols=3)

        if not groups:
            self._append_empty_row(body_table, n_cols=3)
            return

        for g in groups:
            # Заголовок группы (объединённая ячейка через 3 колонки, bold)
            gr = body_table.add_row()
            merged = gr.cells[0].merge(gr.cells[1]).merge(gr.cells[2])
            self._set_cell_text(
                merged,
                str(g.get('title', '')),
                bold=True,
                size=12,
                align='center',
            )

            # Кольца (exterior + holes). Для каждого hole перед точками идёт
            # подзаголовок «Внутренний контур N» в том же формате что и title.
            for ring in g.get('rings', []) or []:
                subtitle = ring.get('subtitle')
                if subtitle:
                    sr = body_table.add_row()
                    merged_s = sr.cells[0].merge(sr.cells[1]).merge(sr.cells[2])
                    self._set_cell_text(
                        merged_s,
                        str(subtitle),
                        bold=False,
                        size=12,
                        align='center',
                    )
                for pt in ring.get('points', []) or []:
                    r = body_table.add_row()
                    cells = r.cells
                    self._set_cell_text(
                        cells[0], str(pt.get('point_num', '-')),
                        size=12, align='center',
                    )
                    self._set_cell_text(
                        cells[1], str(pt.get('x', '-')),
                        size=12, align='center',
                    )
                    self._set_cell_text(
                        cells[2], str(pt.get('y', '-')),
                        size=12, align='center',
                    )

            # Строка S= (только для T5) — выравнивание по левому краю,
            # как в Excel-перечне Fsm_5_3_1 (worksheet.write col 0, align='left').
            if show_area:
                area_sqm = g.get('area_sqm')
                if area_sqm not in (None, '', '0'):
                    ar = body_table.add_row()
                    merged_a = ar.cells[0].merge(ar.cells[1]).merge(ar.cells[2])
                    self._set_cell_text(
                        merged_a,
                        f"S= {area_sqm} кв.м",
                        size=12,
                        align='left',
                    )

    # ------------------------------------------------------------------
    # T7 — ВРИ (3 колонки)
    # ------------------------------------------------------------------
    def _build_table_t7(self, doc, rows: List[Dict[str, Any]]) -> None:
        headers = (
            '№ п/п',
            'Условный номер земельного участка',
            'Вид разрешенного использования',
        )
        header_table, body_table = self._make_table(doc, n_cols=3, widths=self.T7_WIDTHS_CM)
        self._fill_header(header_table, headers)
        self._fill_numbering_row(body_table, n_cols=3)

        if not rows:
            self._append_empty_row(body_table, n_cols=3)
            return

        for row in rows:
            tr = body_table.add_row()
            cells = tr.cells
            self._set_cell_text(cells[0], str(row.get('num', '-')), size=12, align='center')
            self._set_cell_text(cells[1], str(row.get('id_zu', '-')), size=12, align='center')
            self._set_cell_text(cells[2], str(row.get('vri', '-')), size=12, align='left')

    # ------------------------------------------------------------------
    # Table helpers
    # ------------------------------------------------------------------
    def _make_table(self, doc, n_cols: int, widths):
        """Создать ДВЕ separate таблицы для каждого подраздела + paragraph-разделитель.

        - header_table: 1 row column headers (рендерится только в начале)
        - separator paragraph: пустой, height=1 twip (~0.02мм, невидимый)
        - body_table: 1 row numbering + data rows (tblHeader на numbering row →
          ТОЛЬКО numbering повторяется на стр 2+, не column headers)

        КРИТИЧНО: Word merges two adjacent tables в одну logical table если
        между ними нет paragraph. Тогда tblHeader на body_table.row[0] не работает
        (per OOXML §17.4.42: requires contiguous rows from start). Paragraph
        с минимальной height разрывает логическое merging без visual effect.

        Обе таблицы получают fixed layout с одинаковыми column widths →
        колонки header_table визуально сопоставлены с колонками body_table.

        Returns: (header_table, body_table)
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Header table
        header_table = doc.add_table(rows=1, cols=n_cols)
        self.styles.apply_table_borders(header_table, sz=self.styles.TABLE_BORDER_SZ)
        self._apply_fixed_widths(header_table, widths)

        # Paragraph-разделитель: line height = 1 twip (минимум) → invisible
        sep_p = doc.add_paragraph()
        sep_pPr = sep_p._p.get_or_add_pPr()
        sep_spacing = OxmlElement('w:spacing')
        sep_spacing.set(qn('w:line'), '1')          # 1 twip = 0.018мм
        sep_spacing.set(qn('w:lineRule'), 'exact')
        sep_spacing.set(qn('w:before'), '0')
        sep_spacing.set(qn('w:after'), '0')
        sep_pPr.append(sep_spacing)

        # Body table
        body_table = doc.add_table(rows=1, cols=n_cols)
        self.styles.apply_table_borders(body_table, sz=self.styles.TABLE_BORDER_SZ)
        self._apply_fixed_widths(body_table, widths)

        return header_table, body_table

    def _add_after_table_spacer(self, doc) -> None:
        """Пустой абзац-разделитель после таблицы.

        Без него заголовок следующего подраздела рендерится впритык к нижней
        границе таблицы. Параграф пустой, шрифт TNR 12pt — даёт одну строку
        отступа (~5мм при line_spacing=1.5).
        """
        from docx.shared import Pt

        p = doc.add_paragraph()
        run = p.add_run('')
        self.styles.set_run_style(
            run,
            font=self.styles.FONT_NAME,
            size_pt=self.styles.FONT_SIZE_PT,
        )
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    def _apply_fixed_widths(self, table, widths_cm) -> None:
        """Установить фиксированную ширину колонок таблицы (отключить autofit).

        Per OOXML §17.4.49 + Microsoft docs best practice:
        - <w:tblLayout w:type="fixed"/> — отключает autofit (column widths не resize)
        - <w:tblW w:type="dxa" w:w="..."/> — explicit total width в twips
        - <w:tblGrid><w:gridCol w:w="..."/> для каждой колонки — Word использует
          этот grid для layout

        Без этого Word default tblLayout=auto resizes columns based on content,
        и header_table + body_table могут иметь разные column widths visually.
        """
        if not widths_cm:
            return

        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Convert cm widths to twips (1 cm = 567 twips approx, exact via mm_twips)
        widths_twips = [int(round(w_cm * 10 * 56.6929)) for w_cm in widths_cm]
        total_twips = sum(widths_twips)

        # Set per-cell widths (для render compatibility)
        from docx.shared import Cm
        for i, w_cm in enumerate(widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w_cm)

        tbl = table._tbl
        tblPr = tbl.tblPr

        # tblLayout=fixed
        existing_layout = tblPr.find(qn('w:tblLayout'))
        if existing_layout is not None:
            tblPr.remove(existing_layout)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # tblW=dxa explicit total
        existing_w = tblPr.find(qn('w:tblW'))
        if existing_w is not None:
            tblPr.remove(existing_w)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # tblGrid: replace with explicit column widths
        existing_grid = tbl.find(qn('w:tblGrid'))
        if existing_grid is not None:
            tbl.remove(existing_grid)
        tblGrid = OxmlElement('w:tblGrid')
        for w_twips in widths_twips:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(w_twips))
            tblGrid.append(gridCol)
        # Insert tblGrid after tblPr
        tblPr.addnext(tblGrid)

    def _fill_header(self, header_table, headers) -> None:
        """Заполнить header_table row 0 заголовками колонок.

        cantSplit на row 0 (длинные header texts ~50мм auto-height): если не
        помещается на текущей странице, переходит целиком на следующую.
        НЕ tblHeader — header не повторяется на стр 2+.
        """
        row = header_table.rows[0]
        for i, text in enumerate(headers):
            self._set_cell_text(
                row.cells[i], str(text),
                bold=True, size=12, align='center',
            )
        self._set_row_cant_split(row)

    def _fill_numbering_row(self, body_table, n_cols: int) -> None:
        """Заполнить body_table row 0 номерами колонок «1 | 2 | … | N».

        tblHeader=true → numbering повторяется на стр 2+ (per OOXML §17.4.42:
        работает потому что body_table separated от header_table через
        paragraph-разделитель в _make_table).
        """
        row = body_table.rows[0]
        for i in range(n_cols):
            self._set_cell_text(
                row.cells[i], str(i + 1),
                size=12, align='center',
            )
        self._set_row_repeat_header(row)

    def _set_row_cant_split(self, row) -> None:
        """Запретить Word разрезать row across pages — если не вмещается
        в текущей странице, перенести row целиком на следующую."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        existing = trPr.find(qn('w:cantSplit'))
        if existing is None:
            trPr.append(OxmlElement('w:cantSplit'))

    def _set_row_repeat_header(self, row) -> None:
        """Пометить строку как table header — повторяется на каждой новой странице.

        OOXML `<w:tblHeader/>` в `w:trPr`. Word повторяет ТОЛЬКО подряд верхние
        строки от начала таблицы — нельзя пометить только вторую строку без первой.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        for existing in trPr.findall(qn('w:tblHeader')):
            trPr.remove(existing)
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), 'true')
        trPr.append(tblHeader)

    def _append_empty_row(self, table, n_cols: int) -> None:
        """Добавить строку «Отсутствуют» (объединить через все колонки)."""
        tr = table.add_row()
        merged = tr.cells[0]
        for i in range(1, n_cols):
            merged = merged.merge(tr.cells[i])
        self._set_cell_text(merged, 'Отсутствуют', size=12, align='center', italic=True)

    def _set_cell_text(self, cell, text: str, bold: bool = False, italic: bool = False,
                       size: int = 11, align: str = 'left') -> None:
        """Записать текст в ячейку с компактным форматированием для таблиц тела.

        line_spacing=1.0, space_before=0, space_after=0 — иначе ячейки
        раздуваются Word default'ом 8-10pt после абзаца.

        Дублирует Fsm_5_3_9_3._set_cell_text по-намерению (изоляция модулей,
        см. feedback в Phase 5 плана).
        """
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt

        cell.text = ''
        p = cell.paragraphs[0]
        # Компактный paragraph_format — критично для таблиц
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        run = p.add_run(str(text) if text is not None else '')
        self.styles.set_run_style(
            run,
            font=self.styles.FONT_NAME,
            size_pt=size,
            bold=bold,
            italic=italic,
        )
        if align == 'center':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == 'right':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align == 'justify':
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
